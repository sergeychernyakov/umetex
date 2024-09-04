# backend/services/word_translator.py

import os
import logging

import django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'umetex_config.settings')
django.setup()

import docx  # python-docx for manipulating Word documents
from docx.shared import Inches
from typing import Optional, Tuple
from django.conf import settings
from backend.models import Document
from backend.services.text_translator import TextTranslator
from backend.services.yandex_image_translator import YandexImageTranslator
from PIL import Image
from io import BytesIO

logger = logging.getLogger(__name__)

class WordTranslator:
    def __init__(self, document: Document):
        """
        Initialize the WordTranslator with a Document instance.

        :param document: Document instance containing the original Word document and translation details.
        """
        self.document = document
        self.total_pages = 1
        self.current_page = 0
        self.original_word_path = document.original_file.path
        self.translated_file_name = f"translated_{document.pk}.docx"
        self.translations_dir = os.path.join(settings.MEDIA_ROOT, str(document.pk), 'translations')
        os.makedirs(self.translations_dir, exist_ok=True)
        self.translated_file_path = os.path.join(self.translations_dir, self.translated_file_name)
        self.translator = TextTranslator(self.document.translation_language)
        self.yandex_translator = YandexImageTranslator(self.document.translation_language)

    def translate_word(self) -> str:
        """
        Translate the Word document by processing text paragraphs and images, inserting translated content into a new Word document.

        :return: Path to the translated Word document file.
        """
        # Open the original Word document
        original_doc = docx.Document(self.original_word_path)
        translated_doc = docx.Document()

        self.total_pages = max(1, len(original_doc.paragraphs) // 15)

        self.document.update_progress(0, self.total_pages)  # Update progress

        # Collect paragraphs for batch translation
        paragraph_batches = []
        current_batch = []

        for paragraph in original_doc.paragraphs:
            if paragraph.text.strip():
                current_batch.append(paragraph)

            if len(current_batch) == 15:
                paragraph_batches.append(current_batch)
                current_batch = []

        # If there are any remaining paragraphs, add them as a final batch
        if current_batch:
            paragraph_batches.append(current_batch)

        # Process each batch
        for batch in paragraph_batches:
            # Collect all texts in the batch for translation
            texts_to_translate = [paragraph.text for paragraph in batch]
            translated_texts = self.translator.translate_texts(texts_to_translate)  # Batch translate
            # translated_texts = texts_to_translate

            for i, paragraph in enumerate(batch):
                translated_paragraph = translated_doc.add_paragraph()
                translated_text = translated_texts[i]
                translated_run = translated_paragraph.add_run(translated_text)

                # Copy formatting from the original run
                self._copy_paragraph_formatting(paragraph, translated_paragraph)

            # Update progress after processing each batch
            self.current_page = min(self.total_pages - 1, self.current_page + 1)
            self.document.update_progress(self.current_page, self.total_pages)

        # Process images in the Word document
        for rel in original_doc.part.rels.values():
            if "image" in rel.target_ref:
                logger.debug(f"Processing image block in Word document")
                self._process_image_block(rel.target_part.blob, translated_doc)

        # Save the translated Word document
        translated_doc.save(self.translated_file_path)
        self.document.translated_file.name = f'{self.document.pk}/translations/{self.translated_file_name}'
        self.document.save()

        self.document.update_progress(self.total_pages, self.total_pages)

        logger.debug(f"Recreated file saved at: {self.translated_file_path}")
        return self.translated_file_path

    def _process_image_block(self, image_data: bytes, translated_doc: docx.Document):
        """
        Process an image block, translate it, and insert it back into the Word document.

        :param image_data: The image data in bytes.
        :param translated_doc: The new Word document to insert the translated image into.
        """
        # Convert the image data bytes into a PIL Image
        image = Image.open(BytesIO(image_data))

        # Save the image temporarily for translation
        temp_image_path = os.path.join(self.translations_dir, "temp_image.png")
        image.save(temp_image_path)

        # Translate the image using YandexImageTranslator
        translated_image_path = self.yandex_translator.translate_image(image)

        # Insert the translated image back into the Word document
        translated_doc.add_picture(translated_image_path, width=Inches(4))

    def _copy_paragraph_formatting(self, original_paragraph, translated_paragraph):
        """
        Copy formatting from the original paragraph to the translated paragraph.

        :param original_paragraph: The original paragraph object from the source document.
        :param translated_paragraph: The translated paragraph object in the target document.
        """
        for original_run, translated_run in zip(original_paragraph.runs, translated_paragraph.runs):
            translated_run.bold = original_run.bold
            translated_run.italic = original_run.italic
            translated_run.underline = original_run.underline
            translated_run.font.name = original_run.font.name
            translated_run.font.size = original_run.font.size
            translated_run.font.color.rgb = original_run.font.color.rgb

    def __str__(self):
        return self.document.title

# Example usage
# python3 -m backend.services.word_translator
if __name__ == "__main__":
    logger.setLevel(logging.DEBUG)

    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)

    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    console_handler.setFormatter(formatter)

    logger.addHandler(console_handler)

    from backend.models import Document

    # Get a document instance
    document = Document.objects.get(pk=317)  # Replace with actual document ID

    # Create an instance of WordTranslator and translate the document
    word_translator = WordTranslator(document)
    translated_word_path = word_translator.translate_word()
    print(f'Translated Word document saved at: {translated_word_path}')
